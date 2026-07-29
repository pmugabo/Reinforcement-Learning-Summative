"""Build the final evidence-based report from the supplied ALU template."""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path("/workspace/.cache/Machine_Learning_Techniques_II - Summative_Assignment - Report Template.docx")
OUTPUT = ROOT / "report" / "AquaForecast_RL_Summative_Final.docx"

NAVY = "12354D"
BLUE = "167EB2"
PALE = "EAF3F6"
GRAY = "EEF1F3"
HIGHLIGHT = "DDF1E4"


def shade(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_margins(cell, top=50, start=70, bottom=50, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table, color: str = "A8B6BF", size: int = 5) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_cell_text(cell, text: str, size: float = 7.2, bold: bool = False, color: str = "000000", align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], font_size: float = 7.0, best_row: int | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for i, (header, width) in enumerate(zip(headers, widths)):
        table.columns[i].width = Inches(width)
        set_cell_text(table.rows[0].cells[i], header, size=font_size, bold=True, color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(table.rows[0].cells[i], NAVY)
    set_repeat_header(table.rows[0])
    for r_index, values in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(values):
            set_cell_text(cells[i], str(value), size=font_size, align=WD_ALIGN_PARAGRAPH.CENTER)
            if r_index % 2:
                shade(cells[i], "F7F9FA")
            if best_row is not None and r_index + 1 == best_row:
                shade(cells[i], HIGHLIGHT)
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)


def add_body(doc: Document, text: str, size: float = 9.5, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        text = text[len(bold_prefix) :]
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(7.5)


def add_figure(doc: Document, filename: str, width: float, caption: str) -> None:
    picture = doc.add_picture(str(ROOT / "assets" / filename), width=Inches(width))
    picture._inline.docPr.set("title", filename.removesuffix(".png").replace("_", " ").title())
    picture._inline.docPr.set("descr", caption)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def result_text(row: dict[str, str]) -> str:
    return f"{float(row['mean_reward']):.2f} +/- {float(row['std_reward']):.2f}"


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def build() -> None:
    doc = Document(TEMPLATE)
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    section = doc.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_page_number(section.footer.paragraphs[0])

    # Page 1 - template-aligned opening and project overview.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    r = p.add_run("Reinforcement Learning Summative Assignment Report")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_body(doc, "Student Name: Mugabo Patricia", 10.5, "Student Name:")
    add_body(doc, "Video Recording: [ADD VIDEO URL BEFORE SUBMISSION]", 9.5, "Video Recording:")
    add_body(doc, "GitHub Repository: [ADD GITHUB URL BEFORE SUBMISSION]", 9.5, "GitHub Repository:")
    add_heading(doc, "1. Project Overview")
    add_body(doc, "AquaForecast is a mission-based reinforcement-learning extension of a water-demand forecasting and decision-support platform for Kimironko Sector, Gasabo District, Kigali. The simulated agent recommends a daily operating strategy that balances service reliability, reservoir resilience, treatment limits, leakage, energy use and uncertain weather-driven demand. DQN, REINFORCE, PPO and A2C interact with the same stochastic Gymnasium environment for objective comparison. Recommendations remain subject to human operator approval and are not sent directly to infrastructure.")
    add_figure(doc, "environment_visual.png", 6.25, "Figure 1. OpenGL-backed AquaForecast dashboard showing reservoir status, operational observations and the agent recommendation.")

    # Page 2 - agent, start state, actions, terminal conditions.
    doc.add_page_break()
    add_heading(doc, "2. Environment Description")
    add_heading(doc, "2.1 Agent, start state and dynamics", 2)
    add_body(doc, "The agent represents a human-in-the-loop operational planner. Each step equals one day. At reset, the simulator randomizes day of year, 5.4–9.6 ML reservoir storage, recent demand, treatment and pipe health, and 1.2–2.5 ML strategic reserve. Scenario options modify these values for drought, demand surge, high leakage, plant outage or mixed stress. Weather, demand, forecast error, tariff, leakage and equipment health then evolve stochastically.")
    add_heading(doc, "2.2 Discrete action space", 2)
    actions = [
        ["0", "Maintain forecast plan", "Forecast-based production and monitoring"],
        ["1", "Reduce production 20%", "Save energy when shortage risk is low"],
        ["2", "Increase production 20%", "Prepare for moderate peak demand"],
        ["3", "Increase production 40%", "Respond to high demand or low storage"],
        ["4", "Release strategic reserve", "Controlled 0.16 ML release; reserve is finite"],
        ["5", "Activate backup supply", "High-cost 0.13 ML borehole/tanker equivalent"],
        ["6", "Issue conservation advisory", "Five-day voluntary demand reduction"],
        ["7", "Dispatch leak-response crew", "Immediate leak reduction and pipe-health gain"],
        ["8", "Combined drought response", "Production increase plus seven-day advisory"],
        ["9", "Preventive maintenance", "Temporary throughput loss; long-term health gain"],
    ]
    add_table(doc, ["ID", "Action", "Real-world mapping"], actions, [0.42, 2.15, 3.75], font_size=7.6)
    add_heading(doc, "2.3 Terminal conditions", 2)
    add_body(doc, "An episode is truncated after 90 simulated days. It terminates early after three consecutive days with more than 30% unmet demand, representing unacceptable service failure. The tests also cover every action, observation bounds, seeded reset behavior, time-limit behavior and frontend JSON serialization.")

    # Page 3 - observation table and reward.
    doc.add_page_break()
    add_heading(doc, "2.4 Observation space")
    observations = [
        ["Reservoir level", "Stored treated water", "Level/pressure sensor", "float32", "0–12 ML"],
        ["Forecast demand", "Next-day requirement", "AquaForecast ML/API", "float32", "0.20–0.95 ML"],
        ["Previous demand", "Most recent use", "Smart meters/dataset", "float32", "0.25–0.95 ML"],
        ["Rainfall forecast", "Expected rainfall", "Weather API", "float32", "0–25 mm"],
        ["Temperature", "Demand driver", "Weather API/sensor", "float32", "17–35 C"],
        ["Humidity", "Weather context", "Weather API/sensor", "float32", "25–95%"],
        ["Dry-season index", "Seasonal water stress", "Calendar/climate data", "float32", "0–1"],
        ["Leakage risk", "Distribution loss risk", "Flow/pressure sensors", "float32", "2.5–18%"],
        ["Treatment health", "Available plant capacity", "SCADA/maintenance", "float32", "45–100%"],
        ["Energy tariff", "Production cost signal", "Utility API", "float32", "0.40–1.00 norm."],
        ["Forecast uncertainty", "Prediction confidence", "ML model output", "float32", "4–25%"],
        ["Conservation active", "Advisory time remaining", "Campaign database", "float32", "0–7 days"],
        ["Maintenance benefit", "Health benefit remaining", "Maintenance system", "float32", "0–14 days"],
        ["Strategic reserve", "Emergency stock available", "Inventory/level sensor", "float32", "0–2.5 ML"],
        ["Day sine", "Cyclical season encoding", "Calendar", "float32", "0–1"],
        ["Day cosine", "Cyclical season encoding", "Calendar", "float32", "0–1"],
    ]
    add_table(doc, ["Observation", "Meaning", "Real source", "Type", "Physical range"], observations, [1.18, 1.48, 1.65, 0.70, 1.25], font_size=6.45)
    add_heading(doc, "2.5 Reward structure", 2)
    add_body(doc, "The daily reward is R = 5S − 15U² − 4I(U>0.10) − 2|L−0.55| − 1.25E − 6O − C + 0.45B. Here S is service ratio, U unmet-demand ratio, L normalized reservoir level, E normalized production energy cost, O overflow fraction, C action cost including repeat-advisory fatigue, and B indicates safe storage with at least 98.5% service. This prevents unlimited production, emergency supply and repeated conservation messages from becoming cost-free shortcuts.")

    # Page 4 - algorithm design.
    doc.add_page_break()
    add_heading(doc, "3. System Analysis and Design")
    add_heading(doc, "3.1 Shared comparison protocol", 2)
    add_body(doc, "All methods use the same 16-value Box observation, ten-action Discrete space, 90-day horizon, reward function, evaluation seeds and six evaluation scenarios. The policy network uses two 64-unit hidden layers unless the relevant experiment changes hidden width. Best models are selected by mean reward over seeded evaluation episodes, with failure rate used as a safety tie-breaker.")
    add_heading(doc, "3.2 Deep Q-Network (DQN)", 2)
    add_body(doc, "Stable-Baselines3 DQN is the value-based method. It uses an online Q-network, target network, replay buffer, mini-batch temporal-difference updates and epsilon-greedy exploration. Experiments vary learning rate, discount factor, replay capacity, batch size, exploration decay/final epsilon and target update interval.")
    add_heading(doc, "3.3 REINFORCE", 2)
    add_body(doc, "Because Stable-Baselines3 does not ship a REINFORCE class, the required Monte-Carlo policy-gradient algorithm is implemented in PyTorch while retaining the identical Gymnasium environment and evaluation interface. A categorical policy samples actions, discounted returns are normalized, entropy supports exploration and gradients are accumulated across several episodes.")
    add_heading(doc, "3.4 Proximal Policy Optimization (PPO)", 2)
    add_body(doc, "Stable-Baselines3 PPO is an on-policy actor–critic method. It uses generalized advantage estimation, separate policy/value branches and a clipped surrogate objective that limits destabilizing policy updates. Tuning covers learning rate, gamma, rollout length, batch size, GAE lambda, clip range and entropy coefficient.")
    add_heading(doc, "3.5 Advantage Actor–Critic (A2C)", 2)
    add_body(doc, "Stable-Baselines3 A2C jointly optimizes a categorical actor and state-value critic from short on-policy rollouts. Experiments vary learning rate, gamma, rollout length, GAE lambda, entropy coefficient, value-loss weight and maximum gradient norm.")
    add_heading(doc, "3.6 Product integration", 2)
    add_body(doc, "The environment exposes a JSON-serializable payload containing location, timestamp, system measurements, action, explanation and a human-approval notice. This can be returned by an API to an AquaForecast web/mobile dashboard without exposing the training object or granting direct infrastructure control.")

    # Pages 5-6 - complete hyperparameter evidence from the generated summaries.
    dqn = read_csv(ROOT / "logs" / "experiments" / "dqn" / "summary.csv")
    reinforce = read_csv(ROOT / "logs" / "experiments" / "reinforce" / "summary.csv")
    ppo = read_csv(ROOT / "logs" / "experiments" / "ppo" / "summary.csv")
    a2c = read_csv(ROOT / "logs" / "experiments" / "a2c" / "summary.csv")

    doc.add_page_break()
    add_heading(doc, "4. Hyperparameter Experiments")
    add_body(doc, "Each algorithm used ten deliberately different configurations. Stable-Baselines3 methods trained for 25,000 timesteps per run; REINFORCE trained for 250 episodes. Mean reward +/- standard deviation is measured over twelve seeded normal-condition evaluation episodes. Green identifies the promoted model.", 8.4)
    add_heading(doc, "4.1 DQN - 10 runs", 2)
    dqn_rows = [[r["run"], f"{float(r['learning_rate']):.0e}", r["gamma"], r["buffer_size"], r["batch_size"], r["exploration_fraction"], r["exploration_final_eps"], r["target_update_interval"], result_text(r)] for r in dqn]
    add_table(doc, ["Run", "LR", "Gamma", "Buffer", "Batch", "Explore", "Final eps", "Target", "Reward +/- SD"], dqn_rows, [0.31, 0.50, 0.48, 0.60, 0.48, 0.53, 0.55, 0.57, 1.05], font_size=5.45, best_row=8)
    add_body(doc, "DQN run 8 achieved the largest normal-condition mean (420.20), although its SD was 66.29. High gamma (0.999), a 20,000-transition buffer and slower target updates supported long-horizon planning. Run 3 was most stable (SD 5.22), while the smallest learning rate in run 5 learned too slowly and scored 220.49.", 8.1)
    add_heading(doc, "4.2 REINFORCE - 10 runs", 2)
    r_rows = [[r["run"], f"{float(r['learning_rate']):.0e}", r["gamma"], r["hidden_size"], r["entropy_coef"], r["batch_episodes"], result_text(r)] for r in reinforce]
    add_table(doc, ["Run", "LR", "Gamma", "Hidden", "Entropy", "Batch eps", "Reward +/- SD"], r_rows, [0.38, 0.65, 0.58, 0.62, 0.68, 0.72, 1.35], font_size=5.8, best_row=2)
    add_body(doc, "REINFORCE run 2 produced the best mean (400.43) with LR 3e-4, gamma 0.97 and light entropy regularization. Run 10 was slightly lower (392.69) but far more consistent (SD 5.04). Large networks, very high gamma or strong entropy did not improve the 250-episode budget.", 8.1)

    # Page 6 - PPO and A2C.
    doc.add_page_break()
    add_heading(doc, "4.3 PPO - 10 runs", 2)
    ppo_rows = [[r["run"], f"{float(r['learning_rate']):.0e}", r["gamma"], r["n_steps"], r["batch_size"], r["gae_lambda"], r["clip_range"], r["ent_coef"], result_text(r)] for r in ppo]
    add_table(doc, ["Run", "LR", "Gamma", "Steps", "Batch", "GAE", "Clip", "Entropy", "Reward +/- SD"], ppo_rows, [0.31, 0.50, 0.48, 0.50, 0.48, 0.48, 0.48, 0.56, 1.12], font_size=5.45, best_row=3)
    add_body(doc, "PPO run 3 achieved 418.84 +/- 25.41 with zero evaluation failures and only 0.094 ML mean unmet water. The moderate clip (0.20), 256-step rollouts and GAE 0.98 balanced update size and variance. Runs 5 and 8 combined small learning rates with long rollouts and underperformed.", 8.1)
    add_heading(doc, "4.4 A2C - 10 runs", 2)
    a2c_rows = [[r["run"], f"{float(r['learning_rate']):.0e}", r["gamma"], r["n_steps"], r["gae_lambda"], r["ent_coef"], r["vf_coef"], r["max_grad_norm"], result_text(r)] for r in a2c]
    add_table(doc, ["Run", "LR", "Gamma", "Steps", "GAE", "Entropy", "VF coef", "Grad norm", "Reward +/- SD"], a2c_rows, [0.31, 0.50, 0.48, 0.50, 0.48, 0.56, 0.57, 0.59, 1.12], font_size=5.35, best_row=9)
    add_body(doc, "A2C run 9 scored 416.05 with gamma 0.93, ten-step rollouts and the largest tested entropy coefficient (0.04). Run 3 fell to 223.38, indicating sensitivity to the interaction between LR, gamma and rollout length. Because one seed was used per configuration, these are comparative observations rather than isolated causal effects.", 8.1)

    # Page 7 - cumulative performance.
    doc.add_page_break()
    add_heading(doc, "5. Results and Discussion")
    add_heading(doc, "5.1 Cumulative rewards", 2)
    add_figure(doc, "cumulative_rewards_all_methods.png", 5.55, "Figure 2. Cumulative reward for the best run of each method. Episode counts differ slightly because SB3 training stops after the requested timestep budget.")
    add_body(doc, "PPO accumulated 115,866.50 reward across 278 logged episodes and A2C accumulated 115,249.85 across 277. DQN reached 113,845.08 across 280 episodes, while REINFORCE reached 102,379.96 across its fixed 250 episodes. The nearly linear profiles show positive full-horizon returns; slope is fairer than final height when episode counts differ. Mean reward per logged episode was PPO 416.79, A2C 416.06, REINFORCE 409.52 and DQN 406.59.", 8.1)
    add_heading(doc, "5.2 Learning behavior", 2)
    add_figure(doc, "learning_curves.png", 5.35, "Figure 3. Ten-episode rolling-mean reward during training.")
    add_body(doc, "PPO and A2C reached the highest late-training plateau. DQN showed deeper temporary drops, consistent with changing epsilon-greedy data and replay updates. REINFORCE remained smoother but plateaued lower, illustrating a stability-versus-asymptotic-performance trade-off.", 7.9)

    # Page 8 - objective, entropy and convergence.
    doc.add_page_break()
    add_heading(doc, "5.3 Objective and policy entropy", 2)
    add_figure(doc, "training_stability.png", 5.75, "Figure 4. DQN TD-objective loss and policy entropy for the three policy-gradient methods.")
    add_body(doc, "DQN's smoothed TD loss decreased from about 4.17 to 0.33, with later oscillations showing continued replay-based correction rather than divergence. REINFORCE entropy stayed near 2.29 (close to log(10)=2.303), so its policy remained broadly exploratory. PPO entropy declined from about 2.30 to 1.79 and A2C to 2.05, indicating increasingly selective policies without complete entropy collapse.", 8.4)
    add_heading(doc, "5.4 Episodes to converge", 2)
    add_figure(doc, "episodes_to_converge.png", 4.75, "Figure 5. First stable-performance episode; lower is faster.")
    add_body(doc, "Convergence is the first 20-episode rolling mean within 2% of that run's peak for at least eight of the next ten windows. REINFORCE met this rule at episode 25, PPO at 91, A2C at 211 and DQN at 215. REINFORCE therefore stabilized fastest, but its lower final/generalization reward shows that rapid convergence can mean an early plateau rather than the best policy.", 8.3)

    # Page 9 - generalization, conclusion and references.
    doc.add_page_break()
    add_heading(doc, "5.5 Generalization")
    add_figure(doc, "generalization_tests.png", 6.25, "Figure 6. Mean reward across 20 unseen seeds per operational scenario.")
    summary_rows = [
        ["DQN", "420.20 +/- 66.29", "215", "307.04", "0.997", "47.5%"],
        ["REINFORCE", "400.43 +/- 62.36", "25", "294.35", "1.081", "56.7%"],
        ["PPO", "418.84 +/- 25.41", "91", "313.54", "1.268", "38.3%"],
        ["A2C", "416.05 +/- 65.66", "211", "294.31", "1.084", "57.5%"],
    ]
    add_table(doc, ["Method", "Best normal reward", "Conv. ep.", "6-scenario reward", "Unmet ML", "Failure"], summary_rows, [0.86, 1.42, 0.66, 1.10, 0.78, 0.72], font_size=6.6, best_row=3)
    add_body(doc, "PPO had the best six-scenario mean reward (313.54) and lowest average failure rate (38.3%), while DQN had the lowest average unmet water (0.997 ML). Drought and high leakage were handled relatively well; demand surge, plant outage and mixed stress reduced every method's reward. PPO's mixed-stress reward was highest (185.96) but its cumulative unmet water was 3.553 ML because it survived longer, whereas DQN terminated more often and recorded 2.016 ML. This exposes a limitation of total-return comparison when episode lengths vary.", 8.2)
    add_heading(doc, "6. Conclusion and Discussion")
    add_body(doc, "PPO is the recommended demonstration agent because it combined near-best normal reward, lower variability, faster convergence than DQN/A2C, the best average stress-scenario reward and the lowest failure rate. DQN is a strong alternative when minimizing unmet volume is prioritized. REINFORCE was simple and fastest to stabilize but plateaued lower; A2C matched PPO's normal mean less reliably and generalized similarly to REINFORCE. These conclusions are limited by simulated dynamics, one training seed per configuration, reward-weight sensitivity and total reward's dependence on episode length. Future work should use multi-seed tuning, normalize return by days survived, calibrate with utility data, add constrained/safe RL and test multi-zone allocation. All recommendations must remain human-approved before deployment.", 8.2)
    add_heading(doc, "References", 2)
    add_body(doc, "Mnih et al. (2015), Human-level control through deep reinforcement learning. Nature 518. Williams (1992), Simple statistical gradient-following algorithms for connectionist reinforcement learning. Machine Learning 8. Schulman et al. (2017), Proximal Policy Optimization Algorithms. arXiv:1707.06347. Raffin et al. (2021), Stable-Baselines3: Reliable Reinforcement Learning Implementations. JMLR 22.", 7.3)

    doc.core_properties.title = "AquaForecast Mission-Based RL Summative Report"
    doc.core_properties.author = "Mugabo Patricia"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
